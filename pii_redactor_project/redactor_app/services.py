import re
import os
import json
import logging
import spacy
import cv2
import pytesseract
from pytesseract import Output
import fitz  # PyMuPDF
import docx
from docx.shared import RGBColor
from PIL import Image
import requests
from steganography.steganography import Steganography

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# --- Lazy Loading for AI Models ---
transformers = None
torch = None
local_ner_pipeline = None

# --- INITIALIZATION ---
logger = logging.getLogger(__name__)
nlp_spacy = spacy.load("en_core_web_sm")
face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
REDACTION_TEXT = "[REDACTED]"

# --- EXPANDED PII PATTERNS ---
PII_PATTERNS = {
    'EMAIL': re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'),
    'PHONE': re.compile(r'(\+91[\-\s]?)?[0]?(91)?[6-9]\d{9}'),
    'CREDIT_CARD': re.compile(r'\b(?:\d[ -]*?){13,16}\b'),
    'AADHAAR': re.compile(r'\b\d{4}[ -]?\d{4}[ -]?\d{4}\b'),
    'PAN': re.compile(r'[A-Z]{5}[0-9]{4}[A-Z]{1}'),
    'UPI_ID': re.compile(r'[a-zA-Z0-9.\-_]{2,256}@[a-zA-Z]{2,64}'),
    'VOTER_ID': re.compile(r'[A-Z]{3}[0-9]{7}'),
    'DRIVING_LICENCE': re.compile(r'[A-Z]{2}[0-9]{2}\s?[0-9]{4}\s?[0-9]{7}'), # Common DL format
    'GSTIN': re.compile(r'[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z]{1}[1-9A-Z]{1}Z[0-9A-Z]{1}')
}

# --- LOCAL AI MODEL FUNCTION ---
def initialize_local_ner():
    """Initializes the local Hugging Face NER model."""
    global transformers, torch, local_ner_pipeline
    if local_ner_pipeline is None:
        try:
            transformers = __import__('transformers')
            torch = __import__('torch')
            # Switched to a smaller, faster "distilled" model for better CPU performance.
            model_name = "distilbert-base-cased-distilled-squad"
            logger.info(f"Initializing local NER model: {model_name}. This may take a moment...")
            
            local_ner_pipeline = transformers.pipeline(
                "ner", 
                model=model_name, 
                tokenizer=model_name, 
                aggregation_strategy="simple"
            )
            logger.info("Local NER model initialized successfully.")
        except ImportError:
            logger.error("Hugging Face 'transformers' or 'torch' library not found. Please install them.")
            raise
        except Exception as e:
            logger.error(f"Failed to load local NER model: {e}")
            raise

def detect_pii_with_local_ner(text):
    if local_ner_pipeline is None: initialize_local_ner()
    if not text or not text.strip() or not local_ner_pipeline: return []
    try:
        ner_results = local_ner_pipeline(text)
        pii_entities = set()
        for entity in ner_results:
            if entity['entity_group'].upper() in ['PER', 'LOC', 'ORG', 'MISC', 'ANSWER']:
                pii_entities.add(entity['word'])
        logger.info(f"Local model detected {len(pii_entities)} PII elements.")
        return list(pii_entities)
    except Exception as e:
        logger.error(f"Error during local NER processing: {e}")
        return []

def detect_pii_with_gemini(text, api_key):
    if not text or not text.strip(): return []
    headers = {'Content-Type': 'application/json'}
    prompt = f"""Analyze the following text to identify PII. Return ONLY a single, flat JSON array of the exact PII strings found. Example: ["John Doe", "john.doe@email.com"]\n\nText to analyze:\n---\n{text}\n---"""
    payload = json.dumps({"contents": [{"parts": [{"text": prompt}]}]})
    try:
        response = requests.post(f"{GEMINI_API_URL}?key={api_key}", headers=headers, data=payload, timeout=45)
        response.raise_for_status()
        result = response.json()
        content = result['candidates'][0]['content']['parts'][0]['text']
        json_str = content.strip().replace('`', '').lstrip('json').strip()
        return [str(item) for item in json.loads(json_str)]
    except requests.exceptions.Timeout:
        logger.error("Gemini API call timed out.")
        raise ConnectionError("The request to the Gemini API timed out.")
    except (requests.RequestException, KeyError, json.JSONDecodeError) as e:
        logger.error(f"Gemini API call failed: {e}")
        raise ValueError(f"Failed to process response from Gemini API: {e}")


def find_all_pii(text, options):
    all_pii_text = set()
    if options.get('ai_model_choice') == 'local':
        all_pii_text.update(detect_pii_with_local_ner(text))
    elif options.get('gemini_api_key'):
        all_pii_text.update(detect_pii_with_gemini(text, options['gemini_api_key']))

    for pii_type in options.get('pii_types', []):
        if pii_type in PII_PATTERNS:
            for match in PII_PATTERNS[pii_type].finditer(text):
                all_pii_text.add(match.group(0))
    if 'NAME' in options.get('pii_types', []):
        doc = nlp_spacy(text)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                all_pii_text.add(ent.text)
    return sorted(list(all_pii_text), key=len, reverse=True)

def apply_pdf_watermark(page):
    r = page.rect
    text = "REDACTED"
    font_size = min(r.width, r.height) / 10
    text_len = fitz.get_text_length(text, fontname="helv-bold", fontsize=font_size)
    pos = fitz.Point(r.width / 2 - text_len / 2, r.height / 2 + font_size / 4)
    shape = page.new_shape()
    shape.insert_text(pos, text, fontname="helv-bold", fontsize=font_size, color=(0.8, 0.8, 0.8), fill_opacity=0.5, rotate=45)
    shape.commit()

def redact_image(file_path, options):
    output_path = file_path.replace(os.path.splitext(file_path)[1], "_redacted.png")
    img_cv = cv2.cvtColor(cv2.imread(file_path), cv2.COLOR_BGR2RGB)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)
    faces = face_cascade.detectMultiScale(gray, 1.1, 4)
    for (x, y, w, h) in faces:
        cv2.rectangle(img_cv, (x, y), (x + w, y + h), (0, 0, 0), -1)
    ocr_data = pytesseract.image_to_data(img_cv, output_type=Output.DICT)
    full_text = " ".join(ocr_data['text'])
    pii_to_redact = find_all_pii(full_text, options)
    for pii in pii_to_redact:
        for i, text in enumerate(ocr_data['text']):
            if pii in text:
                (x, y, w, h) = (ocr_data['left'][i], ocr_data['top'][i], ocr_data['width'][i], ocr_data['height'][i])
                cv2.rectangle(img_cv, (x, y), (x + w, y + h), (0, 0, 0), -1)
    if options.get('apply_watermark') and not options.get('covert_redaction'):
        h, w, _ = img_cv.shape
        font_scale = min(w, h) / 500
        cv2.putText(img_cv, 'REDACTED', (int(w*0.1), int(h*0.5)), cv2.FONT_HERSHEY_SIMPLEX, font_scale, (128, 128, 128), 2, cv2.LINE_AA)
    final_img_pil = Image.fromarray(img_cv)
    if options.get('covert_redaction'):
        ocr_data_covert = pytesseract.image_to_data(file_path, output_type=Output.DICT)
        full_text_covert = " ".join(ocr_data_covert['text'])
        pii_to_hide = find_all_pii(full_text_covert, options)
        secret_message = "Redacted PII: " + ", ".join(pii_to_hide)
        Steganography.encode(file_path, output_path, secret_message)
        logger.info(f"Covertly hid {len(pii_to_hide)} PII elements in the image.")
        final_img_pil = Image.open(output_path)
    if options.get('wipe_metadata'):
        data = list(final_img_pil.getdata())
        image_without_metadata = Image.new(final_img_pil.mode, final_img_pil.size)
        image_without_metadata.putdata(data)
        image_without_metadata.save(output_path, format='PNG')
    else:
        final_img_pil.save(output_path, format='PNG')
    return output_path

def redact_pdf(file_path, options):
    doc = fitz.open(file_path)
    if options.get('wipe_metadata'): doc.set_metadata({})
    for page in doc:
        text = page.get_text("text")
        pii_to_redact = find_all_pii(text, options)
        for pii in pii_to_redact:
            areas = page.search_for(pii)
            for inst in areas:
                if options.get('covert_redaction'):
                    page.draw_rect(inst, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                else:
                    page.add_redact_annot(inst, fill=(0, 0, 0))
        if not options.get('covert_redaction'): page.apply_redactions()
        if options.get('apply_watermark') and not options.get('covert_redaction'):
            apply_pdf_watermark(page)
    output_path = file_path.replace(".pdf", "_redacted.pdf")
    doc.save(output_path, garbage=4, deflate=True, clean=True)
    doc.close()
    return output_path

def redact_docx(file_path, options):
    doc = docx.Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    pii_to_redact = find_all_pii(full_text, options)
    for pii in pii_to_redact:
        for para in doc.paragraphs:
            if pii in para.text:
                if options.get('covert_redaction'):
                    para.text = para.text.replace(pii, ' ' * len(pii))
                else:
                    para.text = para.text.replace(pii, REDACTION_TEXT)
    output_path = file_path.replace(".docx", "_redacted.docx")
    if options.get('wipe_metadata'):
        doc.core_properties.author = None
        doc.core_properties.last_modified_by = None
    doc.save(output_path)
    return output_path

def process_file(file_path, options):
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()
    if extension in ['.png', '.jpg', '.jpeg']:
        return redact_image(file_path, options)
    elif extension == '.pdf':
        return redact_pdf(file_path, options)
    elif extension == '.docx':
        return redact_docx(file_path, options)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

